#https://platform.openai.com/docs/tutorials/meeting-minutes

import streamlit as st
from docx import Document
from openai import OpenAI
from typing import Literal

client = OpenAI(api_key=st.secrets.openai.api_key)
model1 = "whisper-1"
model2 = "gpt-4o"
model3 = "gpt-3.5-turbo"
sysprompt_abstractsummary = "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
sysprompt_keypoints = "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
sysprompt_actionitetms = "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
sysprompt_sentiment = "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."

class TranscriptionUtilities:
    @staticmethod
    def transcribe_audio(audio_file_path):
        with open(file=audio_file_path, mode="rb") as audio_file:
            transcription = client.audio.transcriptions.create(model=model1, file=audio_file)
            transcription_text = transcription.text
        return transcription_text
    
    @staticmethod
    def get_transcript_extraction(extraction_type: Literal['summary', 'keypoints', 'actionitems', 'sentiment'], transcription: str):
        if extraction_type == "summary":
            sysprompt = sysprompt_abstractsummary
        elif extraction_type == "keypoints":
            sysprompt = sysprompt_keypoints
        elif extraction_type == "actionitems":
            sysprompt = sysprompt_actionitetms
        elif extraction_type == "sentiment":
            sysprompt = sysprompt_sentiment
        else:
            sysprompt = sysprompt_abstractsummary
        
        messages = [{"role": "system", "content": sysprompt}, {"role": "user", "content": transcription}]
        response = client.chat.completions.create(messages=messages, model=model2)
        response_text = response.choices[0].message.content
        return response_text
    
    @staticmethod
    def meeting_minutes(transcription):
        summary = TranscriptionUtilities.abstract_summary_extraction(transcription=transcription)
        keypoints = TranscriptionUtilities.key_points_extraction(transcription=transcription)
        actionitems = TranscriptionUtilities.action_items_extraction(transcription=transcription)
        sentiment = TranscriptionUtilities.sentiment_extraction(transcription=transcription)
        meetingminutes = {"abstract_summary": summary, "key_points": keypoints, "action_items": actionitems, "sentiment": sentiment}
        return meetingminutes
    
    @staticmethod
    def abstract_summary_extraction(transcription):
        messages = [{"role": "system", "content": sysprompt_abstractsummary}, {"role": "user", "content": transcription}]
        response = client.chat.completions.create(messages=messages, model=model2)
        response_text = response.choices[0].message.content
        return response_text
    
    @staticmethod
    def key_points_extraction(transcription):
        messages = [{"role": "system", "content": sysprompt_keypoints}, {"role": "user", "content": transcription}]
        response = client.chat.completions.create(messages=messages, model=model2)
        response_text = response.choices[0].message.content
        return response_text
    
    @staticmethod
    def action_items_extraction(transcription):
        messages = [{"role": "system", "content": sysprompt_actionitetms}, {"role": "user", "content": transcription}]
        response = client.chat.completions.create(messages=messages, model=model2)
        response_text = response.choices[0].message.content
        return response_text
    
    @staticmethod
    def sentiment_extraction(transcription):
        messages = [{"role": "system", "content": sysprompt_sentiment}, {"role": "user", "content": transcription}]
        response = client.chat.completions.create(messages=messages, model=model2)
        response_text = response.choices[0].message.content
        return response_text

    @staticmethod
    def save_as_doc(minutes, filename):
        doc = Document()
        for key, value in minutes.items():
            heading = ' '.join(word.capitalize() for word in key.split('_'))
            doc.add_heading(heading, level=1)
            doc.add_paragraph(value)
            # Add a line break between sections
            doc.add_paragraph()
        doc.save(filename)

# audio_file_path = "Earningscall.wav"
# transcription = TranscriptionUtilities.transcribe_audio(audio_file_path)
# minutes = TranscriptionUtilities.meeting_minutes(transcription)
# print(minutes)

# TranscriptionUtilities.save_as_docx(minutes, 'meeting_minutes.docx')

class MeetingMinutes:
    def __init__(self, audio_file_path):
        self.audio_file_path = audio_file_path
        self.transcription = TranscriptionUtilities.transcribe_audio(audio_file_path=self.audio_file_path)
        self.minutes = TranscriptionUtilities.meeting_minutes(transcription=self.transcription)
        self.doc = TranscriptionUtilities.save_as_doc(minutes=self.minutes, filename="meeting_minutes.docx")

transcription = TranscriptionUtilities.transcribe_audio(audio_file_path="EarningsCall.wav")
print(transcription)